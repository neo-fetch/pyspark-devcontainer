from docx import Document
import re
import time

document = Document("IT.docx")


def prompt():
    full_name = input("Enter full name: ")
    serial_number = input("Enter serial number: ")

    laptop_model = input("Enter laptop model: ")
    laptop_serial_number = input("Enter laptop serial number (Host Name): ")
    laptop_ram_size = input("Enter laptop RAM size (RAM): ")
    laptop_hard_drive_size = input("Enter laptop hard drive size (Hard Disk): ")
    laptop_screen_size = input("Enter laptop screen size (Screen): ")

    validity_date = input("Enter validity date (dd/mm/yyyy): ")
    # ensure Validity date is in the format dd/mm/yyyy
    while not re.match(
        r"^(?:(?:31(\/|-|\.)(?:0?[13578]|1[02]))\1|(?:(?:29|30)(\/|-|\.)(?:0?[1,3-9]|1[0-2])\2))(?:(?:1[6-9]|[2-9]\d)?\d{2})$|^(?:29(\/|-|\.)0?2\3(?:(?:(?:1[6-9]|[2-9]\d)?(?:0[48]|[2468][048]|[13579][26])|(?:(?:16|[2468][048]|[3579][26])00))))$|^(?:0?[1-9]|1\d|2[0-8])(\/|-|\.)(?:(?:0?[1-9])|(?:1[0-2]))\4(?:(?:1[6-9]|[2-9]\d)?\d{2})$",
        validity_date
    ):
        validity_date = input("Invalid date format. Enter validity date (dd/mm/yyyy): ")
    
    passport_number = input("Enter passport number: ")
    passport_number = passport_number.strip()

    auth_name = input("Enter name of person authorizing: ")
    today = time.strftime("%d/%m/%Y")
    date_of_issue = input(f"Enter date of issue (dd/mm/yyyy): Default is today's date: {today}")
    if date_of_issue == "":
        date_of_issue = today
    # Return a tuple of the values
    return (full_name, serial_number, laptop_model, laptop_serial_number, laptop_ram_size, laptop_hard_drive_size, laptop_screen_size, validity_date, passport_number, auth_name, date_of_issue)

def replace_text(text, new_text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            # preserve the formatting of the paragraph
            runs_font = [run.font.name for run in paragraph.runs]
            runs_size = [run.font.size for run in paragraph.runs]
            # runs_bold = [run.bold for run in paragraph.runs]
            paragraph.text = paragraph.text.replace(text, new_text)
            for i in range(len(paragraph.runs)):
                paragraph.runs[i].font.name = runs_font[i]
                paragraph.runs[i].font.size = runs_size[i]
                # Make it bold too
                # paragraph.runs[i].bold = runs_bold[i]

def replace_table_text(text, new_text):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style = paragraph.style.name
                    if text in paragraph.text:
                        paragraph.text = paragraph.text.replace(text, new_text)
                        paragraph.style.name = style

def fill_document():
    # Get the values from the prompt
    full_name, serial_number, laptop_model, laptop_serial_number, laptop_ram_size, laptop_hard_drive_size, laptop_screen_size, validity_date, passport_number, auth_name, date_of_issue = prompt()
    # Replace the text in the document
    replace_text("<Full Name>", full_name)
    replace_text("<Serial Number>", serial_number)
    replace_text("<Laptop Model>", laptop_model)
    replace_text("<Laptop Number>", laptop_serial_number)
    replace_text("<RAM Size>", laptop_ram_size)
    replace_text("<Storage Size>", laptop_hard_drive_size)
    replace_text("<Screen Size>", laptop_screen_size)
    replace_text("<Validity Date>", validity_date)
    replace_text("<Passport Number>", passport_number)
    replace_text("<Auth Name>", auth_name)
    replace_text("<Current Date>", date_of_issue)
    
    # Save the document
    document.save('IT_filled.docx')

if __name__ == "__main__":
    fill_document()